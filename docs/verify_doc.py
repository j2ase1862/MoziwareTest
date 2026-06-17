# -*- coding: utf-8 -*-
import os
from docx import Document
p = r"D:\BODA_글라스_PoC_소개서_매뉴얼.docx"
doc = Document(p)
print("이미지(inline):", len(doc.inline_shapes))
print("표:", len(doc.tables))
print("H1 섹션:")
for pp in doc.paragraphs:
    if pp.style.name == "Heading 1":
        print("  -", pp.text)
print("아키텍처 PNG:", os.path.exists(r"D:\Repo\Android\doc-architecture.png"))
print("파일 크기:", os.path.getsize(p), "bytes")
