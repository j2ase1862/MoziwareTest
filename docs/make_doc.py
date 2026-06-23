# -*- coding: utf-8 -*-
"""BODA 글라스 + BODA.VMS.Web PoC 소개서/매뉴얼 (.docx) 생성 + 아키텍처 다이어그램."""
import os
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = r"D:\Repo\Android"
OUT  = r"D:\BODA_글라스_PoC_소개서_매뉴얼.docx"
ARCH = os.path.join(BASE, "doc-architecture.png")
FONT = r"C:\Windows\Fonts\malgun.ttf"
FONTB = r"C:\Windows\Fonts\malgunbd.ttf"

# ---------------------------------------------------------------- 아키텍처 다이어그램
def make_arch():
    W, H = 1240, 660
    img = Image.new("RGB", (W, H), "#FFFFFF")
    d = ImageDraw.Draw(img)
    def f(sz, bold=False):
        return ImageFont.truetype(FONTB if bold else FONT, sz)
    def box(x, y, w, h, title, sub, fill, border, tcol="#0A0C10", dash=False):
        d.rounded_rectangle([x, y, x+w, y+h], radius=14, fill=fill, outline=border, width=3)
        d.text((x+w/2, y+h/2-16), title, font=f(22, True), fill=tcol, anchor="mm")
        if sub:
            d.text((x+w/2, y+h/2+14), sub, font=f(15), fill="#334155", anchor="mm")
    def arrow(x1, y1, x2, y2, label=""):
        d.line([x1, y1, x2, y2], fill="#64748B", width=3)
        # arrowhead
        import math
        ang = math.atan2(y2-y1, x2-x1)
        for s in (-0.4, 0.4):
            d.line([x2, y2, x2-14*math.cos(ang-s), y2-14*math.sin(ang-s)], fill="#64748B", width=3)
        if label:
            d.text(((x1+x2)/2, (y1+y2)/2-14), label, font=f(14), fill="#1E40AF", anchor="mm")

    d.text((40, 28), "시스템 구성도", font=f(26, True), fill="#0A0C10")

    # 행: 글라스 → Cloudflare → 서버 → DB
    y = 250
    box(40, y, 250, 110, "스마트 글라스", "Moziware Cimo · BODA 글라스 앱", "#DBEAFE", "#3B82F6")
    box(360, y, 220, 110, "게이트웨이", "PoC: Cloudflare · 운영: 고객 환경", "#FEF3C7", "#F59E0B")
    box(650, y, 250, 110, "BODA 서버", "PoC: boda-vms.com · 운영: 고객 인프라", "#DCFCE7", "#22C55E")
    box(970, y, 230, 110, "데이터베이스", "SQLite · WarehouseItem/OutboundOrder", "#F1F5F9", "#64748B")
    arrow(290, y+55, 360, y+55, "REST/HTTPS")
    arrow(580, y+55, 650, y+55, "")
    arrow(900, y+55, 970, y+55, "EF Core")

    # 웹 관리자 (상단) — SignalR
    box(650, 70, 250, 90, "웹 관리자(브라우저)", "입고/출고 관리 · 실시간", "#EDE9FE", "#8B5CF6")
    arrow(775, 160, 775, y, "")
    arrow(775, y, 775, 162, "SignalR 실시간")

    # 고객 ERP/WMS (하단, 향후)
    box(650, 470, 250, 90, "고객 ERP / WMS", "주문·재고 연동 (운영)", "#FFFFFF", "#94A3B8", tcol="#475569")
    d.line([775, 360, 775, 470], fill="#CBD5E1", width=3)
    d.text((790, 415), "연동/맞춤개발", font=f(13), fill="#94A3B8", anchor="lm")

    d.text((40, H-34),
           "※ Cloudflare·boda-vms.com은 PoC 데모 환경입니다. 운영은 고객사 인프라(온프렘/클라우드)와 "
           "ERP·WMS에 맞춰 구성·연동·맞춤개발합니다.",
           font=f(14), fill="#64748B")
    img.save(ARCH)

# ---------------------------------------------------------------- docx helpers
def set_korean_font(doc):
    for sname in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        try:
            st = doc.styles[sname]
            st.font.name = "맑은 고딕"
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), "맑은 고딕")
            rfonts.set(qn("w:ascii"), "맑은 고딕")
            rfonts.set(qn("w:hAnsi"), "맑은 고딕")
        except KeyError:
            pass

def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(doc, text, size=11, bold=False, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(it).font.size = Pt(11)

def image(doc, path, width_cm=14.5, caption=None):
    if not os.path.exists(path):
        para(doc, f"[이미지 없음: {os.path.basename(path)}]", size=9, color=(0xB0,0x30,0x30))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run("▲ " + caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x5D, 0x6B)
        r.italic = True

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return t

# ---------------------------------------------------------------- 문서 작성
def img(name):
    return os.path.join(BASE, name)

def build():
    make_arch()
    doc = Document()
    set_korean_font(doc)
    sec = doc.sections[0]
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)

    # ===== 표지 =====
    for _ in range(3):
        doc.add_paragraph()
    para(doc, "BODA 글라스  ×  BODA.VMS.Web", size=26, bold=True,
         color=(0x1E,0x40,0xAF), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "스마트 글라스 입출고 관리 솔루션", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "제품 소개서 & 사용 매뉴얼", size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55,0x5D,0x6B), space_after=20)
    image(doc, img("screenshot-newcimo.png"), width_cm=12,
          caption="BODA 글라스 메인 화면 (Moziware Cimo, 854×480)")
    doc.add_paragraph()
    para(doc, "고객 PoC 자료 · 2026-06-17", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55,0x5D,0x6B))
    doc.add_page_break()

    # ===== 1. 개요 =====
    h(doc, "1. 개요", 1)
    para(doc, "BODA 글라스는 단안 스마트 글라스(Moziware Cimo)에서 바코드 스캔과 음성 명령만으로 "
              "물류 입고 위치를 조회하고 출고 주문을 핸즈프리로 피킹하는 네이티브 Android 앱입니다. "
              "글라스는 화면·음성·스캔만 담당하는 얇은 클라이언트로, 데이터/로직은 백엔드 서버가 처리하고 "
              "작업 결과는 웹 관리 화면에 실시간으로 반영됩니다.")
    para(doc, "※ 본 PoC는 데모 환경(서버 boda-vms.com, Cloudflare)에서 검증한 것이며, 실제 운영은 "
              "고객사 인프라(온프렘/클라우드)와 ERP·WMS에 맞춰 구성·연동·맞춤개발됩니다.",
         size=10, color=(0x55,0x5D,0x6B))
    para(doc, "핵심 가치", bold=True, space_after=2)
    bullets(doc, [
        "핸즈프리: 두 손이 자유로운 상태로 음성/내장 스캐너만으로 작업",
        "현장 즉시성: 입고 위치를 대형 고대비 화면에 즉시 표시, 출고는 위치로 직접 안내",
        "검증된 피킹: 제품 바코드 스캔으로 오피킹 방지 + 수량 카운트 + 출하 확정",
        "유연한 연동: 신규 서버 없이 백엔드 확장으로 구현 — 고객사 VMS/ERP/WMS에 맞춰 연동·맞춤개발 가능",
        "실시간 가시성: 글라스 작업이 웹 관리 대시보드에 새로고침 없이 즉시 반영",
    ])

    # ===== 2. 시스템 구성 =====
    h(doc, "2. 시스템 구성", 1)
    image(doc, ARCH, width_cm=16, caption="전체 시스템 구성도")
    para(doc, "구성 요소", bold=True, space_after=2)
    table(doc, ["구성 요소", "역할/기술"], [
        ["스마트 글라스", "Moziware Cimo (RealWear 기반, Android 10) · 내장 바코드 스캐너 · WearHF 음성"],
        ["BODA 글라스 앱", "네이티브 Android(Kotlin) · 얇은 클라이언트(화면·음성·스캔만)"],
        ["백엔드 서버", "ASP.NET Core + Blazor · PoC: boda-vms.com / 운영: 고객 인프라 배포 또는 고객 시스템 연동"],
        ["게이트웨이", "PoC: Cloudflare(HTTPS/Tunnel) / 운영: 고객 네트워크·보안 정책에 맞춤"],
        ["데이터베이스", "PoC: SQLite · 운영: 고객 표준 DB / ERP·WMS 데이터 연동"],
        ["실시간", "SignalR — 글라스 작업 → 웹 관리 화면 즉시 반영"],
    ])
    para(doc, "데이터 흐름", bold=True, space_after=2)
    bullets(doc, [
        "글라스 → 서버: HTTPS REST 호출(/api/glass/...) — 익명 + X-API-Key 보호",
        "서버 → DB: 입고 위치 조회 / 출고 주문·피킹 상태 저장",
        "서버 → 웹 관리자: SignalR로 출고 상태 변경 실시간 푸시",
        "(운영) 고객 ERP/WMS ↔ 서버: 주문·재고 연동(맞춤개발)",
    ])
    para(doc, "※ PoC 환경 안내: boda-vms.com·Cloudflare·SQLite는 시연용 데모 구성입니다. "
              "운영 도입 시 백엔드는 고객사 인프라(온프렘/클라우드)에 배포하거나 고객사 ERP/WMS와 연동하며, "
              "네트워크·보안·데이터 표준은 고객 환경에 맞춰 구성합니다.",
         size=10, color=(0x55,0x5D,0x6B))

    # ===== 3. 주요 기능 =====
    h(doc, "3. 주요 기능", 1)
    table(doc, ["기능", "설명"], [
        ["입고 위치 조회", "제품 바코드 스캔 → '입고/입고제품' → 보관 위치를 대형 표시"],
        ["입고 위치 등록", "(웹) 바코드별 위치(Zone-Rack-Level-Bin)·좌표 등록/수정"],
        ["출고 목록", "(글라스) 활성 출고 주문을 보고 선택 — 커서/번호/탭"],
        ["출고 피킹", "주문 선택 → 라인별 위치 안내 → 제품 스캔 검증 + 수량 → 출하 확정"],
        ["출고 오더 관리", "(웹) 출고 주문·라인·목적지 등록/수정, 진행 상태 확인"],
        ["실시간 반영", "글라스 피킹/출하 확정이 웹 관리 화면에 즉시 표시"],
    ])

    # ===== 4. 스마트 글라스 앱 매뉴얼 =====
    h(doc, "4. 스마트 글라스 앱(BODA 글라스) 매뉴얼", 1)

    h(doc, "4.1 설치 및 실행", 2)
    bullets(doc, [
        "APK를 기기에 사이드로드(설치) — 기기 저장소에 영구 설치(USB 분리·재부팅 후에도 유지)",
        "런처(앱 목록)에서 'BODA 글라스' 아이콘 실행, 또는 음성으로 앱 실행",
        "Wi-Fi 인터넷이 boda-vms.com(443)에 연결되면 어디서든 사용 가능",
    ])
    image(doc, img("screenshot-icon.png"), width_cm=13, caption="앱 아이콘 'BODA 글라스' (앱 정보 화면)")

    h(doc, "4.2 메인 화면", 2)
    image(doc, img("screenshot-newcimo.png"), width_cm=14.5, caption="메인 화면 — 버튼 텍스트가 곧 음성 명령")
    para(doc, "버튼(=음성 명령)", bold=True, space_after=2)
    table(doc, ["명령", "동작"], [
        ["바코드 스캔", "내장 스캐너로 제품 바코드 스캔"],
        ["입고 / 입고제품", "스캔한 바코드의 입고(보관) 위치 조회·표시"],
        ["출고", "출고 목록 화면으로 이동"],
        ["처음으로", "대기 상태로 초기화"],
        ["종료", "앱 완전 종료"],
    ])

    h(doc, "4.3 입고 위치 조회", 2)
    image(doc, img("screenshot-inbound-map.png"), width_cm=14.5,
          caption="입고 위치 조회 — 2D 미니맵(구역/랙 핀) + 대형 위치 코드 + 방향 힌트")
    para(doc, "① '바코드 스캔' → 제품 바코드 스캔  ② '입고' 또는 '입고제품'  ③ 보관 위치가 대형 고대비로 표시되고, "
              "왼쪽 2D 미니맵에 구역/랙 핀과 '○구역 · ○랙 · ○단 · ○칸' 힌트가 함께 표시됩니다. "
              "등록되지 않은 바코드는 빨간색으로 '등록되지 않은 바코드'를 안내합니다.")

    h(doc, "4.4 출고 — 주문 목록", 2)
    image(doc, img("doc-orderlist.png"), width_cm=14.5, caption="출고 목록 — 활성 주문 선택")
    para(doc, "주문 선택 3가지 방법(병행)", bold=True, space_after=2)
    table(doc, ["방법", "사용법"], [
        ["커서(권장)", "'다음'/'이전'으로 강조 이동 → '열기'"],
        ["번호", "각 행 배지 번호로 '항목 1 열기'처럼 발화"],
        ["탭", "행을 직접 터치"],
        ["주문 스캔", "주문/송장 바코드를 스캔해 바로 진입"],
    ])

    h(doc, "4.5 출고 — 피킹(스캔 검증)", 2)
    image(doc, img("screenshot-minimap.png"), width_cm=14.5,
          caption="피킹 화면 — 2D 미니맵·위치·품목·수량 / 라인별 제품 스캔 검증")
    bullets(doc, [
        "각 라인의 '위치'로 이동 → 그 자리에서 '스캔'으로 제품 바코드 스캔",
        "맞는 품목이면 수량이 1씩 증가(예: 0/3 → 3/3), 다 채우면 자동으로 다음 라인",
        "다른 품목을 스캔하면 빨간색 '이 주문에 없는 품목'으로 차단(오피킹 방지)",
        "'이전'/'다음'으로 라인 이동, 상단 '취소'로 언제든 홈 복귀",
    ])

    h(doc, "4.6 출고 — 출하 확정", 2)
    image(doc, img("screenshot-pick-dest.png"), width_cm=14.5,
          caption="출하 확정 화면 — 전 라인 피킹 시 '출하 확정' 활성, 출고 목적지 표시")
    para(doc, "전 라인 피킹이 끝나면 '출하 확정' 버튼이 활성화되고, 확정 시 주문 상태가 '완료(Done)'로 바뀌며 "
              "출고 목적지(도크/출고대)가 표시됩니다. 미완료 상태에서는 확정이 비활성으로 보호됩니다.")

    h(doc, "4.7 음성 명령 빠른 참조", 2)
    table(doc, ["화면", "음성 명령"], [
        ["메인", "바코드 스캔 · 입고 · 입고제품 · 출고 · 처음으로 · 종료"],
        ["출고 목록", "이전 · 다음 · 열기 · 항목 N 열기 · 주문 스캔 · 새로고침 · 취소"],
        ["피킹", "스캔 · 이전 · 다음 · 출하 확정 · 닫기 · 취소"],
        ["전역(WearHF)", "뒤로 가기 · 명령어"],
    ])

    # ===== 5. 웹 관리 =====
    h(doc, "5. 웹 관리 (BODA.VMS.Web)", 1)
    para(doc, "boda-vms.com에 로그인 후, 좌측 메뉴에서 입출고 마스터 데이터를 관리합니다. "
              "글라스의 작업은 이 화면에 실시간으로 반영됩니다.")
    h(doc, "5.1 입고 위치 등록", 2)
    bullets(doc, [
        "메뉴: '입고 위치 등록'",
        "바코드 · 품목코드/명 · 위치(Zone-Rack-Level-Bin) · 3D 좌표 등록/수정/삭제",
        "여기 등록된 바코드를 글라스에서 스캔하면 즉시 위치가 조회됨",
    ])
    h(doc, "5.2 출고 오더 관리", 2)
    bullets(doc, [
        "메뉴: '출고 피킹'",
        "출고 주문(주문번호·고객·배송지·출고 목적지) + 피킹 라인(바코드·수량) 등록/수정",
        "상태(대기 → 피킹중 → 완료)가 글라스 작업에 따라 실시간 갱신",
    ])
    h(doc, "5.3 실시간(SignalR)", 2)
    para(doc, "글라스에서 제품을 스캔(피킹)하거나 출하 확정하면, 열려 있는 웹 관리 화면의 해당 주문 상태가 "
              "새로고침 없이 '대기 → 피킹중 → 완료'로 즉시 바뀝니다. 운영자는 현장 진행 상황을 실시간으로 모니터링할 수 있습니다.")
    para(doc, "※ 웹 관리 화면 캡처는 시연 환경(로그인 계정)에서 추가 첨부 예정.", size=9, color=(0x55,0x5D,0x6B))

    # ===== 6. PoC 시연 시나리오 =====
    h(doc, "6. PoC 시연 시나리오", 1)
    para(doc, "준비물", bold=True, space_after=2)
    bullets(doc, [
        "BODA 글라스 설치된 Moziware Cimo (Wi-Fi 연결)",
        "시연용 주문 QR / 제품 바코드(QR), 웹 관리 화면(브라우저, 로그인)",
    ])
    para(doc, "시연 순서", bold=True, space_after=2)
    table(doc, ["단계", "내용"], [
        ["① 입고 조회", "제품 바코드 스캔 → '입고' → 보관 위치 대형 표시"],
        ["② 출고 진입", "'출고' → 출고 목록에서 주문 선택(커서/번호/탭) 또는 '주문 스캔'"],
        ["③ 피킹", "라인별 위치 안내 → 제품 스캔으로 수량 채움(오피킹 차단 시연)"],
        ["④ 출하 확정", "전 라인 완료 → '출하 확정' → 목적지 표시(상태 '완료')"],
        ["⑤ 실시간 확인", "웹 '출고 피킹' 화면이 새로고침 없이 '피킹중 → 완료'로 변하는지 확인"],
    ])
    para(doc, "시연용 주문 QR 예시 (SO-2026-001)", bold=True, space_after=2)
    image(doc, img("qr-SO-2026-001.png"), width_cm=4.5, caption="주문 QR: SO-2026-001 (가나상사, 3개 라인)")
    para(doc, "시드 데이터(예시)", bold=True, space_after=2)
    table(doc, ["주문", "고객", "목적지", "라인"], [
        ["SO-2026-001", "가나상사", "3번 도크 / 출고대 B", "P-1001×3, P-1003×2, P-1007×1"],
        ["SO-2026-002", "대한물산", "1번 도크 / 출고대 A", "P-1002×5, P-1010×4"],
    ])

    # ===== 7. 운영/도입 요건 =====
    h(doc, "7. 운영 / 도입 요건", 1)
    bullets(doc, [
        "기기: Moziware Cimo(Android 10) — 내장 바코드 스캐너 / WearHF 음성",
        "네트워크: PoC는 Wi-Fi로 boda-vms.com(HTTPS) 접속. 운영은 고객 네트워크 정책에 맞춰 "
        "사내망(온프렘) 또는 고객 클라우드로 구성(앱 서버 주소만 교체)",
        "백엔드 배포: PoC는 데모 서버. 운영은 고객 인프라에 배포하거나 고객 기존 시스템(VMS/ERP/WMS)과 연동",
        "보안: 핸즈프리 익명 + X-API-Key(운영 시 키 강제). HTTPS 종단·게이트웨이는 고객 보안 정책에 맞춤",
        "데이터 출처: 입고 위치·출고 주문 공급(고객 ERP/WMS 연동·CSV·수기) 방식 확정이 도입 1순위 과제",
    ])

    # ===== 8. 기능 범위 (가능/조건부/불가능) =====
    h(doc, "8. 기능 범위 — 가능 / 조건부 / 불가능", 1)
    para(doc, "고객 기대치를 정확히 맞추기 위해, 현재 PoC에서 동작하는 범위 · 추가 데이터/개발이 필요한 범위 · "
              "기기(단안 글라스) 한계로 불가능한 범위를 구분합니다.")

    h(doc, "8.1 가능 — PoC에서 동작 (✅)", 2)
    bullets(doc, [
        "바코드/QR 스캔 → 입고 위치 즉시 조회(대형 표시 + 2D 미니맵 + 방향 힌트)",
        "음성·핸즈프리 조작(화면 버튼 텍스트가 곧 음성 명령)",
        "출고 피킹: 주문 목록 선택(음성/탭) → 위치 안내 → 제품 스캔 검증 + 수량 카운트 → 출하 확정",
        "오피킹 방지(잘못된 품목 스캔 차단) · 라인별 진행률",
        "웹 관리(입고 위치·출고 오더 등록/수정) + 실시간(SignalR) 상태 반영",
        "앱 내 위치 시각화 — 구역/랙 도식 미니맵",
    ])

    h(doc, "8.2 조건부 가능 — 추가 데이터/개발 필요 (△)", 2)
    bullets(doc, [
        "실제 창고 배치도 미니맵: 2D 평면도 이미지(또는 구역/랙 좌표) + 각 위치 실측 좌표 필요",
        "고객 ERP/WMS 연동(주문·재고 동기화): 연동 인터페이스 개발",
        "체크인(웨이포인트) 길안내: 통로/랙에 QR 부착 + 앱 보강 — 추가 HW 없이, 스캔 시점마다 다음 위치 안내",
        "재고 차감 · 작업자 배정(내 작업) · 피킹 리스트 인쇄/출고 리포트",
        "한국어 음성 인식 정확도: 명령어·소음 환경 튜닝(짧은 단어 → 또렷한 문구)",
    ])

    h(doc, "8.3 현재 불가능 — 기기(단안 HUD) 한계 (✕)", 2)
    bullets(doc, [
        "실시간 연속 실내 길안내(이동에 따른 위치 추적): 실내는 GPS 불가 + 단안 HUD라 공간추적(SLAM) 없음 "
        "→ BLE/UWB 등 실내 측위 인프라 구축 시에만 가능(별도 프로젝트)",
        "XR/AR 오버레이(실제 선반에 화살표·하이라이트를 공간 고정): 공간 AR 기기(HoloLens/Magic Leap 등) 필요 "
        "— Moziware Cimo(단안 보조 디스플레이)로는 불가",
        "몰입형 VR · 정밀 3D 뷰: 단안 · 좁은 FOV(20°)로 가독성이 떨어져 부적합",
    ])
    para(doc, "※ 위 '불가능' 항목은 Moziware Cimo(단안 HUD) 기준입니다. 공간 AR 기기/실내 측위 인프라를 "
              "도입하면 일부는 향후 구현 가능합니다.", size=9, color=(0x55,0x5D,0x6B))

    # ===== 9. 향후 확장 =====
    h(doc, "9. 향후 확장 로드맵", 1)
    bullets(doc, [
        "고객 ERP/WMS 연동: 판매오더·재고 자동 동기화(시드/수기 → 운영 데이터)",
        "재고 차감: 출하 확정 시 보관 위치 재고 반영",
        "작업자 배정: 로그인/배정 기반 '내 작업' 목록",
        "2차 목표(3D): 입고/피킹 위치를 공장 도면에 3D 하이라이트",
        "라벨/리포트: 피킹 리스트 인쇄, 출고 실적 리포트",
    ])

    doc.save(OUT)
    print("저장:", OUT)

build()
